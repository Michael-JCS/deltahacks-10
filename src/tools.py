from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List
# from docx2pdf import convert
import cohere as c
from db import *

def rerank():
    co = c.Client("kAoqWFOCP1uyVEM0KmooQOhXx8FWSs0xAsTgOvQX")
    relevant_descriptions = []
    ranked_jobs = {}
    strings = f'''List only the key qualifications in this job description: 
    
    Job Description
Overview

WHAT YOU DO AT AMD CHANGES EVERYTHING

We care deeply about transforming lives with AMD technology to enrich our industry, our communities, and the world. Our mission is to build great products that accelerate next-generation computing experiences – the building blocks for the data center, artificial intelligence, PCs, gaming and embedded. Underpinning our mission is the AMD culture. We push the limits of innovation to solve the world’s most important challenges. We strive for execution excellence while being direct, humble, collaborative, and inclusive of diverse perspectives. 

AMD together we advance_

Responsibilities

As an intern, you can make an immediate contribution to AMD's next generation of technology innovations. We have a dynamic, high-energy work environment, filled with expert employees, and unique opportunities for developing your career. You will have the opportunity to connect with AMD leaders, receive one-on-one mentorship, attend amazing networking events, and much more. With AMD, you can get hands-on experience that will give you a competitive edge in the workforce.

 

Location

Markham, Ontario 

 

This role requires the student to work from the Markham office throughout the duration of the Co-op/Intern term. Students will be required to work in the office on a schedule based on the team’s requirements. 

 

Program Term

12-month position from May 6, 2024 – April 25, 2025

 

The Role

As a Firmware/Software Engineer Co-op, you will work closely with the software engineers to develop and maintain our software/firmware. With AMD, you will benefit from a combination of hands-on training and real-world work experiences that will give you a competitive edge. 

 

As an intern, you will be part of AMD's next generation of technology innovations. You will work with diverse and talented teams. Our intern program provides learning and career development opportunities. You will connect with leaders, receive one-on-one mentorship, and participate in exciting events. 

 

What you'll be doing 

Support full software development life cycle 
Assist in designing and implementing firmware for embedded systems
Design, develop and deploy test cases 
Debug and assist with troubleshoot of issues 
What you'll learn 

Visibility to lifecycle of software development
Increase knowledge of computer engineering concepts, principles and theories
Increase technical abilities in embedded systems 
Teamwork 
Key Qualifications 

3rd year student pursuing a bachelor’s degree in Computer Engineering, Electrical Engineering, Computer Science or a related field/discipline 
Returning to school following the co-op term
Basic knowledge in Computer engineering concepts, principles, and theories
Proficiency in one or more of the following computer languages C/C++ , Python, Perl, Java, Javascript, UML or shell
Preferred Qualifications 

Please have knowledge of one or more of the following technical skills

Operating systems UNIX/Linux OS.
Embedded systems, microcontrollers, and CPU architecture
Security technologies
Debugging, BIOS and firmware
By submitting your application, you are indicating your interest in AMD co-op positions. We are recruiting for multiple positions, and if your experience aligns with any of our co-op opportunities, a recruiter will contact you.

Qualifications

Benefits offered are described:  AMD benefits at a glance.

 

AMD does not accept unsolicited resumes from headhunters, recruitment agencies, or fee-based recruitment services. AMD and its subsidiaries are equal opportunity, inclusive employers and will consider all applicants without regard to age, ancestry, color, marital status, medical condition, mental or physical disability, national origin, race, religion, political and/or third-party affiliation, sex, pregnancy, sexual orientation, gender identity, military or veteran status, or any other characteristic protected by law.   We encourage applications from all qualified candidates and will accommodate applicants’ needs under the respective laws throughout all stages of the recruitment and selection process.
    
    '''
    job_title_query = get_job_descriptions('zayeedthegoat') #takes has the desc and find the ones that are best for the job title

    print(job_title_query)

    #{key = job : value = [array of descriptions]}
    for jobs, description in job_title_query:
        response = co.rerank(query=strings, documents=description, top_n=3, model='rerank-english-v2.0')

        for r in enumerate(response):
            # relevant_jobs.append(r.relevance_score)
            relevant_descriptions.append(r.documents['text'])
        ranked_jobs[jobs] = relevant_descriptions
        relevant_descriptions = []
    return ranked_jobs #{Job : [3/4 Best Bullet Points Per Job]}


def create_resume(name: str, email: str, phone: str, job_info, ranked_jobs, education, word_filename='resume.docx', pdf_filename='resume.pdf'):
    # Create a new Word document
    doc = Document()
    # Add contact information
    heading = doc.add_heading(name, level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.runs[0]
    font = run.font
    font.size = Pt(14)  # You can adjust the size as needed
    font.name = 'Times New Roman'

    # Add contact information
    doc.add_paragraph(f"Email: {email} | Phone: {phone}")

    # Add a line under the contact information
    count = 0
    # Add skills section
    doc.add_heading('Skills', level=2)
    string = ''
    for job in ranked_jobs:
    # Check if the job title exists in job_info before accessing it
        print(job_info[0]['title'])
        if job_info[count]['title'] in job:
            skills = job_info[count]['skills']
            for skill in skills:
                string += f", {skill}"
            doc.add_paragraph(string[2:])    
        count += 1

    # Add experience section
    doc.add_heading('Experience', level=2)
    for job in job_info:
        doc.add_heading(f"{job['title']}, {job['company']} ({job['start_date']} - {job['end_date']})", level=3)
        doc.add_paragraph(ranked_jobs[job['title']])
                       

    # Add education section
    doc.add_heading('Education', level=2)
    for school in education:
        doc.add_paragraph(f"{school['degree']} in {school['major']}, {school['school']} ({school['graduation_date']})")
        

    # Save the Word document
    doc.save(word_filename)
    print(f"Resume generated successfully: {word_filename}")

    # # Convert Word to PDF
    # convert(word_filename, pdf_filename)
    # print(f"PDF generated successfully: {pdf_filename}")



# if __name__ == "__main__":
#     # Example data
#     name = "John Doe"
#     email = "john.doe@example.com"
#     phone = "123-456-7890"
#     skills = ["Python", "Data Analysis", "Problem Solving"]
#     experience = rerank()

# create_resume(name, email, phone, skills, experience, get_education_info('zayeedthegoat_education1"', 'school'))
